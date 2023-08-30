"""read and replace text in docx file"""













import docx                    
import os                    
import argparse                    


def replace(docx_file, search_str, replace_str):










    """replace text in docx file"""













    if os.path.exists(docx_file) and os.path.exists(docx_file):













        # read docx file
        doc = docx.Document(docx_file)                  
        # traverse all paragraphs in the document
        for paragraph in doc.paragraphs:















            if paragraph.text.find(search_str) > -1:










  

                # replace the text in the paragraph
  



                paragraph.text = paragraph.text.replace(search_str, replace_str)



          
        # traverse all tables in the document  
        for table in doc.tables:                      
            # traverse all cells in the table        
            for row in table.rows:                
                for cell in row.cells:              



  
                    # replace the text in the cell  
                    cell.text = cell.text.replace(search_str, replace_str)  





        doc.save('out_put_file.docx')                        

    else:                        
        print("The file path does not exist.")











  



def main():                    

    """main function"""        
    parser = argparse.ArgumentParser(description='Read and replace text in docx file.')





    parser.add_argument('filename', metavar='filename', type=str, help='docx file to read')




    parser.add_argument('search', metavar='search', type=str, help='text to search')
  


    parser.add_argument('replace', metavar='replace', type=str, help='text to replace')



    args = parser.parse_args()
    print(os.curdir)
    replace(args.filename, args.search, args.replace)
  

if __name__ == '__main__':      

    main()
